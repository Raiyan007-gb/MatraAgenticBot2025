import io
import os
import datetime
import tempfile
from PIL import Image
import fitz  # PyMuPDF
import markdown2
import re
import textwrap
from docx import Document
from docx.shared import Inches # For controlling image size

def resize_logo(logo_path: str) -> str:
    """Resize the logo PNG to 2496x2505 pixels, preserving aspect ratio with transparent padding."""
    try:
        # Open the logo image
        img = Image.open(logo_path).convert("RGBA")
        
        # Calculate scaling to fit within 2496x2505 while preserving aspect ratio
        target_width, target_height = 2496, 2505
        img_ratio = img.width / img.height
        target_ratio = target_width / target_height
        
        if img_ratio > target_ratio:
            # Image is wider; scale by width
            new_width = target_width
            new_height = int(target_width / img_ratio)
        else:
            # Image is taller; scale by height
            new_height = target_height
            new_width = int(target_height * img_ratio)
        
        # Resize image with high-quality LANCZOS filter
        img_resized = img.resize((new_width, new_height), Image.LANCZOS)
        
        # Create a new blank image with transparent background
        final_img = Image.new("RGBA", (target_width, target_height), (255, 255, 255, 0))
        
        # Paste resized image in the center
        paste_x = (target_width - new_width) // 2
        paste_y = (target_height - new_height) // 2
        final_img.paste(img_resized, (paste_x, paste_y))
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        final_img.save(temp_file.name, format="PNG")
        temp_file.close()
        return temp_file.name
    except Exception as e:
        print(f"Error resizing logo: {e}")
        return logo_path  # Fallback to original logo if resizing fails

def generate_docx(markdown_content: str, logo_path: str = None) -> io.BytesIO:
    doc = Document()

    # Add logo if provided
    if logo_path:
        try:
            # Resize logo to 2496x2505
            resized_logo_path = resize_logo(logo_path)
            # Add logo at the beginning, adjust size as needed
            doc.add_picture(resized_logo_path, width=Inches(1.5))
            # Clean up temporary file
            if resized_logo_path != logo_path and os.path.exists(resized_logo_path):
                os.unlink(resized_logo_path)
        except Exception as e:
            print(f"Warning: Could not add logo to DOCX: {e}")
    
    # Basic Markdown to DOCX conversion
    # This is a simplified parser. For complex markdown, you might need a more robust solution
    # like converting markdown to HTML first (with markdown2) and then parsing HTML,
    # or using pandoc via a subprocess.

    # Split content into lines for processing
    lines = markdown_content.split('\n')
    
    current_paragraph_text = []

    for line in lines:
        stripped_line = line.strip()
        
        # Handle Headings
        if stripped_line.startswith('### '): # H3
            if current_paragraph_text: doc.add_paragraph(" ".join(current_paragraph_text)); current_paragraph_text = []
            doc.add_heading(stripped_line[4:], level=3)
        elif stripped_line.startswith('## '): # H2
            if current_paragraph_text: doc.add_paragraph(" ".join(current_paragraph_text)); current_paragraph_text = []
            doc.add_heading(stripped_line[3:], level=2)
        elif stripped_line.startswith('# '): # H1
            if current_paragraph_text: doc.add_paragraph(" ".join(current_paragraph_text)); current_paragraph_text = []
            doc.add_heading(stripped_line[2:], level=1)
        
        # Handle Lists (very basic: assumes lines starting with '*' or '-' are list items)
        elif stripped_line.startswith(('* ', '- ')):
            if current_paragraph_text: doc.add_paragraph(" ".join(current_paragraph_text)); current_paragraph_text = []
            # python-docx uses styles for bullet points. 'ListBullet' is a common one.
            doc.add_paragraph(stripped_line[2:], style='ListBullet')
        
        # Handle empty lines (interpreted as paragraph breaks)
        elif not stripped_line:
            if current_paragraph_text:
                doc.add_paragraph(" ".join(current_paragraph_text))
                current_paragraph_text = []
        
        # Regular text lines
        else:
            current_paragraph_text.append(stripped_line)

    # Add any remaining text as a final paragraph
    if current_paragraph_text:
        doc.add_paragraph(" ".join(current_paragraph_text))
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_transparent_watermark(logo_path, opacity=20):
    """Create a transparent watermark from a logo."""
    try:
        img = Image.open(logo_path).convert("RGBA")
        watermark = Image.new('RGBA', img.size, (255, 255, 255, 0))
        data = img.getdata()
        new_data = []
        for item in data:
            if item[3] > 0:
                new_data.append((item[0], item[1], item[2], int(opacity * 255 / 100)))
            else:
                new_data.append(item)
        watermark.putdata(new_data)
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        watermark.save(temp_file.name)
        temp_file.close()
        return temp_file.name
    except Exception as e:
        print(f"Error creating watermark: {e}")
        return None

def parse_policy_sections(policy_md):
    """Parse markdown into sections with headings and content."""
    sections = []
    current_heading = None
    current_content = []
    
    for line in policy_md.split('\n'):
        if line.startswith('## '):
            if current_heading:
                sections.append({'heading': current_heading, 'content': '\n'.join(current_content).strip()})
            current_heading = line[3:].strip()
            current_content = []
        else:
            current_content.append(line)
    
    if current_heading:
        sections.append({'heading': current_heading, 'content': '\n'.join(current_content).strip()})
    
    return sections

def extract_subcategory(heading):
    """Extract subcategory from heading."""
    return heading.split(':')[0].strip() if ':' in heading else heading.strip()

def add_footer_to_page(page, page_number, static_footer, current_sections, margin_left, margin_right, page_width, page_height, margin_bottom, margin_top):
    """Helper function to add footer to a page with dynamic section titles."""
    footer_y = page_height - margin_bottom + 10  # Adjusted for tighter spacing
    
    # Add decorative line above footer
    page.draw_line(
        fitz.Point(margin_left, footer_y - 8),
        fitz.Point(page_width - margin_right, footer_y - 8),
        color=(0.2, 0.4, 0.7),  # Blue
        width=0.5
    )
    
    # Create section titles text for footer
    section_info = ""
    if current_sections:
        unique_sections = list(dict.fromkeys(extract_subcategory(section) for section in current_sections))  # Use subcategory
        section_info = f"Sections: {', '.join(unique_sections)}"
    
    # Combine footer text components
    footer_text = f"{static_footer}\nAdapted from NIST AI RMF\n{section_info}"
    
    # Define footer rectangle with adjusted height for smaller font
    footer_rect = fitz.Rect(
        margin_left,
        footer_y,
        page_width - margin_right,
        page_height - 15
    )
    
    # Insert footer text with smaller font and tighter line spacing
    page.insert_textbox(
        footer_rect,
        footer_text,
        fontname="helvetica",
        fontsize=7,  # Smaller font size
        align=fitz.TEXT_ALIGN_CENTER,
        lineheight=1.1  # Tighter line spacing
    )
    
    # Add page number at top-right corner
    page_num_rect = fitz.Rect(
        page_width - margin_right - 40,
        margin_top - 20,
        page_width - margin_right,
        margin_top
    )
    
    page.insert_textbox(
        page_num_rect,
        f"Page {page_number}",
        fontname="helvetica",
        fontsize=7,  # Smaller font size
        align=fitz.TEXT_ALIGN_RIGHT
    )

def generate_pdf(policy_md, logo_path=None):
    """Generate a PDF document from markdown content with continuous sections flow."""
    doc = fitz.open()
    sections = parse_policy_sections(policy_md)
    current_year = datetime.datetime.now().year
    static_footer = f"© NIST AI RMF {current_year} | All Rights Reserved | Do Not Use Without Permission"
    
    # Create watermark if logo provided
    watermark_path = None
    if logo_path:
        watermark_path = create_transparent_watermark(logo_path)
    
    # Define page dimensions and margins
    page_width = 595  # A4 width in points
    page_height = 842  # A4 height in points
    margin_left = 72
    margin_right = 72
    margin_top = 72
    margin_bottom = 72
    
    # Create cover page
    cover_page = doc.new_page(width=page_width, height=page_height)
    
    # Add watermark to cover page if provided
    if watermark_path:
        try:
            img = Image.open(watermark_path)
            img_width = page_width * 0.6
            img_height = img_width * (img.height / img.width)
            x = (page_width - img_width) / 2
            y = (page_height - img_height) / 2
            cover_page.insert_image(
                fitz.Rect(x, y, x + img_width, y + img_height),
                filename=watermark_path,
                keep_proportion=True
            )
        except Exception as e:
            print(f"Error inserting watermark: {e}")
    
    # Add logo to cover page if provided
    if logo_path:
        logo_y_pos = 150
        try:
            # Resize logo to 2496x2505
            resized_logo_path = resize_logo(logo_path)
            img = Image.open(resized_logo_path)
            aspect_ratio = img.width / img.height
            logo_width = 120
            logo_height = 120 / aspect_ratio
            cover_page.insert_image(
                fitz.Rect(
                    (page_width - logo_width) / 2,
                    logo_y_pos,
                    (page_width + logo_width) / 2,
                    logo_y_pos + logo_height
                ),
                filename=resized_logo_path
            )
            # Clean up temporary file
            if resized_logo_path != logo_path and os.path.exists(resized_logo_path):
                os.unlink(resized_logo_path)
        except Exception as e:
            print(f"Error inserting logo: {e}")
    
    # Add title to cover page
    title_y_pos = 300
    cover_page.insert_textbox(
        fitz.Rect(
            margin_left,
            title_y_pos,
            page_width - margin_right,
            title_y_pos + 100
        ),
        "NIST AI Risk Management Framework",
        fontname="helvetica-bold",
        fontsize=24,
        color=(0.2, 0.4, 0.7),  # Blue
        align=fitz.TEXT_ALIGN_CENTER
    )
    
    # Add date to cover page
    subtitle_y_pos = title_y_pos + 60
    cover_page.insert_textbox(
        fitz.Rect(
            margin_left,
            subtitle_y_pos,
            page_width - margin_right,
            subtitle_y_pos + 50
        ),
        f"Generated on {datetime.datetime.now().strftime('%B %d, %Y')}",
        fontname="helvetica",
        fontsize=14,
        color=(0.3, 0.3, 0.3),  # Dark gray
        align=fitz.TEXT_ALIGN_CENTER
    )
    
    # Initialize variables for content flow
    current_page = doc.new_page(width=page_width, height=page_height)
    current_page_number = 2  # Starting after cover page
    y_pos = margin_top
    current_page_sections = []  # Track sections on current page
    
    # Add watermark and logo to first content page
    if watermark_path:
        try:
            img = Image.open(watermark_path)
            img_width = page_width * 0.6
            img_height = img_width * (img.height / img.width)
            x = (page_width - img_width) / 2
            y = (page_height - img_height) / 2
            current_page.insert_image(
                fitz.Rect(x, y, x + img_width, y + img_height),
                filename=watermark_path,
                keep_proportion=True
            )
        except Exception as e:
            print(f"Error inserting watermark: {e}")
    
    if logo_path:
        try:
            # Resize logo to 2496x2505
            resized_logo_path = resize_logo(logo_path)
            header_rect = fitz.Rect(
                page_width - margin_right - 890,
                margin_top - 50,  # Adjusted to move logo higher
                page_width - margin_right,
                margin_top - 20   # Adjusted to maintain logo height
            )
            current_page.insert_image(
                header_rect,
                filename=resized_logo_path,
                keep_proportion=True
            )
            # Clean up temporary file
            if resized_logo_path != logo_path and os.path.exists(resized_logo_path):
                os.unlink(resized_logo_path)
        except Exception as e:
            print(f"Error inserting header logo: {e}")
    
    # Process each section
    for section_data in sections:
        section_title = section_data['heading']  # Use full heading for footer
        
        # Check if we need a new page before adding section heading
        if y_pos > page_height - margin_bottom - 80 and y_pos != margin_top:
            add_footer_to_page(current_page, current_page_number, static_footer, current_page_sections,
                              margin_left, margin_right, page_width, page_height, margin_bottom, margin_top)
            current_page = doc.new_page(width=page_width, height=page_height)
            current_page_number += 1
            y_pos = margin_top
            current_page_sections = [section_title]  # Start new page with current section
            
            # Add watermark and logo to new page
            if watermark_path:
                try:
                    img = Image.open(watermark_path)
                    img_width = page_width * 0.6
                    img_height = img_width * (img.height / img.width)
                    x = (page_width - img_width) / 2
                    y = (page_height - img_height) / 2
                    current_page.insert_image(
                        fitz.Rect(x, y, x + img_width, y + img_height),
                        filename=watermark_path,
                        keep_proportion=True
                    )
                except Exception as e:
                    print(f"Error inserting watermark: {e}")
            if logo_path:
                try:
                    # Resize logo to 2496x2505
                    resized_logo_path = resize_logo(logo_path)
                    header_rect = fitz.Rect(
                        page_width - margin_right - 890,
                        margin_top - 50,  # Adjusted to move logo higher
                        page_width - margin_right,
                        margin_top - 20   # Adjusted to maintain logo height
                    )
                    current_page.insert_image(
                        header_rect,
                        filename=resized_logo_path,
                        keep_proportion=True
                    )
                    # Clean up temporary file
                    if resized_logo_path != logo_path and os.path.exists(resized_logo_path):
                        os.unlink(resized_logo_path)
                except Exception as e:
                    print(f"Error inserting header logo: {e}")
        
        # Add section heading
        heading_bg_rect = fitz.Rect(
            margin_left - 5,
            y_pos - 5,
            page_width - margin_right + 5,
            y_pos + 25
        )
        current_page.draw_rect(heading_bg_rect, fill=(0.9, 0.95, 1.0), width=0)
        
        heading_rect = fitz.Rect(
            margin_left,
            y_pos,
            page_width - margin_right,
            y_pos + 30
        )
        current_page.insert_textbox(
            heading_rect,
            section_data['heading'],
            fontname="helvetica-bold",
            fontsize=16,
            color=(0.2, 0.4, 0.7),
            align=fitz.TEXT_ALIGN_LEFT
        )
        y_pos += 40
        if section_title not in current_page_sections:
            current_page_sections.append(section_title)  # Add section title to current page
        
        # Convert markdown to HTML
        html = markdown2.markdown(section_data['content'], extras=["fenced-code-blocks", "tables"])
        
        # Parse and render content
        paragraphs = re.split(r'<p>|</p>|<br>|<br/>', html)
        for para in paragraphs:
            if not para.strip():
                continue
            
            # Check if we need a new page
            if y_pos > page_height - margin_bottom - 50:
                add_footer_to_page(current_page, current_page_number, static_footer, current_page_sections,
                                  margin_left, margin_right, page_width, page_height, margin_bottom, margin_top)
                current_page = doc.new_page(width=page_width, height=page_height)
                current_page_number += 1
                y_pos = margin_top
                current_page_sections = [section_title]  # Continue same section on new page
                
                # Add watermark and logo to new page
                if watermark_path:
                    try:
                        img = Image.open(watermark_path)
                        img_width = page_width * 0.6
                        img_height = img_width * (img.height / img.width)
                        x = (page_width - img_width) / 2
                        y = (page_height - img_height) / 2
                        current_page.insert_image(
                            fitz.Rect(x, y, x + img_width, y + img_height),
                            filename=watermark_path,
                            keep_proportion=True
                        )
                    except Exception as e:
                        print(f"Error inserting watermark: {e}")
                if logo_path:
                    try:
                        # Resize logo to 2496x2505
                        resized_logo_path = resize_logo(logo_path)
                        header_rect = fitz.Rect(
                            page_width - margin_right - 890,
                            margin_top - 50,  # Adjusted to move logo higher
                            page_width - margin_right,
                            margin_top - 20   # Adjusted to maintain logo height
                        )
                        current_page.insert_image(
                            header_rect,
                            filename=resized_logo_path,
                            keep_proportion=True
                        )
                        # Clean up temporary file
                        if resized_logo_path != logo_path and os.path.exists(resized_logo_path):
                            os.unlink(resized_logo_path)
                    except Exception as e:
                        print(f"Error inserting header logo: {e}")
            
            # Handle lists with improved parsing
            if re.search(r'<[uo]l>.*?</[uo]l>', para, re.DOTALL):
                list_items = re.findall(r'<li>(.*?)</li>', para, re.DOTALL | re.MULTILINE)
                for item in list_items:
                    if item.strip():
                        clean_item = re.sub(r'<[^>]+>', '', item).strip()
                        if clean_item:
                            # Check if we need a new page for this list item
                            if y_pos > page_height - margin_bottom - 20:
                                add_footer_to_page(current_page, current_page_number, static_footer, current_page_sections,
                                                  margin_left, margin_right, page_width, page_height, margin_bottom, margin_top)
                                current_page = doc.new_page(width=page_width, height=page_height)
                                current_page_number += 1
                                y_pos = margin_top
                                current_page_sections = [section_title]  # Continue same section
                                
                                # Add watermark and logo to new page
                                if watermark_path:
                                    try:
                                        img = Image.open(watermark_path)
                                        img_width = page_width * 0.6
                                        img_height = img_width * (img.height / img.width)
                                        x = (page_width - img_width) / 2
                                        y = (page_height - img_height) / 2
                                        current_page.insert_image(
                                            fitz.Rect(x, y, x + img_width, y + img_height),
                                            filename=watermark_path,
                                            keep_proportion=True
                                        )
                                    except Exception as e:
                                        print(f"Error inserting watermark: {e}")
                                if logo_path:
                                    try:
                                        # Resize logo to 2496x2505
                                        resized_logo_path = resize_logo(logo_path)
                                        header_rect = fitz.Rect(
                                            page_width - margin_right - 890,
                                            margin_top - 50,  # Adjusted to move logo higher
                                            page_width - margin_right,
                                            margin_top - 20   # Adjusted to maintain logo height
                                        )
                                        current_page.insert_image(
                                            header_rect,
                                            filename=resized_logo_path,
                                            keep_proportion=True
                                        )
                                        # Clean up temporary file
                                        if resized_logo_path != logo_path and os.path.exists(resized_logo_path):
                                            os.unlink(resized_logo_path)
                                    except Exception as e:
                                        print(f"Error inserting header logo: {e}")
                            
                            # Add bullet point
                            current_page.insert_text(
                                fitz.Point(margin_left + 10, y_pos + 11),
                                "•",
                                fontname="times-roman",
                                fontsize=11
                            )
                            
                            # Add indented text with adjusted height
                            rect = fitz.Rect(
                                margin_left + 25,
                                y_pos,
                                page_width - margin_right,
                                y_pos + 50
                            )
                            text_height = current_page.insert_textbox(
                                rect,
                                clean_item,
                                fontname="times-roman",
                                fontsize=11,
                                align=fitz.TEXT_ALIGN_LEFT
                            )
                            wrapped_lines = textwrap.wrap(clean_item, width=80)
                            y_pos += max(20, len(wrapped_lines) * 15) + 5
                continue
            
            # Handle regular paragraphs
            clean_para = re.sub(r'<[^>]+>', '', para).strip()
            if clean_para:
                rect = fitz.Rect(
                    margin_left,
                    y_pos,
                    page_width - margin_right,
                    y_pos + 100
                )
                text_length = current_page.insert_textbox(
                    rect,
                    clean_para,
                    fontname="times-roman",
                    fontsize=11,
                    align=fitz.TEXT_ALIGN_LEFT
                )
                wrapped_lines = textwrap.wrap(clean_para, width=80)
                y_pos += max(20, len(wrapped_lines) * 15) + 10
    
    # Add footer to the last page
    add_footer_to_page(current_page, current_page_number, static_footer, current_page_sections,
                      margin_left, margin_right, page_width, page_height, margin_bottom, margin_top)
    
    # Clean up temporary watermark file if created
    if watermark_path and os.path.exists(watermark_path):
        try:
            os.unlink(watermark_path)
        except:
            pass
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# def main():
#     try:
#         # Define file paths
#         markdown_file_path = os.path.join(os.path.dirname(__file__), 'documents', 'test.md')
#         logo_path = '/Users/rajin/Downloads/Data_Island.png'
        
#         # Read markdown content
#         with open(markdown_file_path, 'r', encoding='utf-8') as f:
#             policy_md = f.read()
        
#         # Validate logo is a PNG and exists
#         if not logo_path.lower().endswith('.png'):
#             raise ValueError("Logo must be a PNG file.")
#         if not os.path.exists(logo_path):
#             raise FileNotFoundError(f"Logo file not found: {logo_path}")
        
#         # Validate markdown file exists
#         if not os.path.exists(markdown_file_path):
#             raise FileNotFoundError(f"Markdown file not found: {markdown_file_path}")
        
#         # Generate PDF
#         pdf_buffer = generate_pdf(policy_md, logo_path)
        
#         # Save PDF to the same directory as the script
#         output_dir = os.path.dirname(os.path.abspath(__file__))
#         output_path = os.path.join(output_dir, 'policy.pdf')
        
#         with open(output_path, 'wb') as f:
#             f.write(pdf_buffer.getvalue())
        
#         print(f"PDF generated successfully at: {output_path}")
        
#     except Exception as e:
#         print(f"Error generating PDF: {str(e)}")

# if __name__ == "__main__":
#     main()